"""生成项目技术文档 Word 文件"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

# ====== 样式设置 ======
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    heading = doc.styles[f'Heading {i}']
    heading.font.name = '微软雅黑'
    heading.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

# ====== 封面 ======
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('建筑行业 AI 智能助手')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('基于 RAG 的土木工程规范问答系统')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run('技术文档 v1.0')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

for _ in range(6):
    doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'编写日期：{datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(11)

info2 = doc.add_paragraph()
info2.alignment = WD_ALIGN_PARAGRAPH.CENTER
info2.add_run('技术栈：Python 3.10+ / FastAPI / React 18 / LlamaIndex / DeepSeek').font.size = Pt(10)

doc.add_page_break()

# ====== 目录 ======
doc.add_heading('目录', level=1)
toc_items = [
    '1. 项目概述',
    '2. 技术栈',
    '3. 系统架构',
    '    3.1 整体数据流',
    '    3.2 用户交互路径',
    '    3.3 索引单例架构',
    '    3.4 MCP 服务器与 Agent 架构',
    '4. 项目结构',
    '    4.1 后端模块详表',
    '    4.2 前端组件详表',
    '5. 环境变量与配置',
    '6. API 参考',
    '    6.1 健康检查',
    '    6.2 RAG 问答（流式 / 非流式）',
    '    6.3 Agent 问答',
    '    6.4 文件管理',
    '7. 安装与运行',
    '    7.1 环境要求',
    '    7.2 快速启动',
    '    7.3 数据管道',
    '8. 关键实现细节',
    '9. 架构约束',
    '10. 部署说明',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ====== 1. 项目概述 ======
doc.add_heading('1. 项目概述', level=1)
doc.add_paragraph(
    '本项目是一个基于 RAG（检索增强生成）技术的建筑行业 AI 智能问答系统。'
    '以中国国家标准 (GB)、行业标准 (JGJ/JTG/SL) 和工程建设法律法规为知识库，'
    '为建筑工程从业人员提供专业规范查询与智能解答。'
)
doc.add_paragraph('核心功能：')
features = [
    '规范知识库：收录 200+ 项土木工程标准规范，231 个文档文件，覆盖结构设计、施工验收、安全规范、材料标准等类别',
    'RAG 检索增强：LlamaIndex 框架 + 本地 BAAI/bge-small-zh-v1.5 Embedding + ModelScope Rerank 重排序',
    '深度推理模式：USE_REASONING=true 时启用 DeepSeek Reasoner，原生思维链真流式输出',
    '用户文件上传：支持 PDF/Word/PPT/Excel/TXT/Markdown，独立 FAISS 向量索引，问答时优先检索',
    '流式响应：SSE (Server-Sent Events) 实时逐字输出，支持 reasoning（思考过程）和 token（回答）双通道',
    'Agent 智能搜索：CrewAI Agent + RAG 知识库 + 百度搜索 + Tavily 搜索三工具协作',
    'MCP 服务器：基于 FastMCP 框架，注册 baidu_search 和 tavily_search 工具，供外部 MCP 客户端调用',
    'Slate 设计系统：Inter 字体 + Slate 色系，专业简洁的 React 前端界面',
]
for f in features:
    doc.add_paragraph(f'  - {f}')

# ====== 2. 技术栈 ======
doc.add_heading('2. 技术栈', level=1)
table = doc.add_table(rows=11, cols=2, style='Light Grid Accent 1')
cells = [
    ('层级', '技术'),
    ('LLM', 'DeepSeek (deepseek-chat / deepseek-reasoner)，OpenAI 兼容 API'),
    ('RAG 框架', 'LlamaIndex (llama-index-core >= 0.11.0)'),
    ('Embedding', '本地 BAAI/bge-small-zh-v1.5 (sentence-transformers) / API Qwen/Qwen3-Embedding-8B'),
    ('Rerank', 'Qwen/Qwen3-Reranker-8B via ModelScope API（失败自动降级）'),
    ('用户文件向量库', 'FAISS (IndexFlatIP + IndexIDMap，L2 归一化，内积 = 余弦相似度)'),
    ('后端', 'FastAPI + uvicorn (Python 3.10+)'),
    ('前端', 'React 18 + Vite 5，marked 库渲染 GFM，DOMPurify 消毒'),
    ('Agent 框架', 'CrewAI (Agent + Task + Crew + 自定义 BaseTool)'),
    ('MCP', 'FastMCP（服务器）+ MCP Python SDK（客户端/stdio 子进程桥接）'),
    ('文件解析', 'PDF: PyMuPDF / Word: python-docx / PPT: python-pptx / Excel: openpyxl'),
]
for i, (col1, col2) in enumerate(cells):
    table.cell(i, 0).text = col1
    table.cell(i, 1).text = col2
    for cell in table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                run.font.name = '微软雅黑'

doc.add_paragraph()

# ====== 3. 系统架构 ======
doc.add_heading('3. 系统架构', level=1)

doc.add_heading('3.1 整体数据流', level=2)
doc.add_paragraph(
    '文档加载流程：文档 (.txt/.pdf/.docx/.pptx/.xlsx/.md) '
    '-> 统一文件解析器 file_parser.py -> SentenceSplitter 中文切分 -> '
    'Embedding 向量化 -> VectorStoreIndex（持久化到 backend/storage/）'
)
doc.add_paragraph(
    '问答流程：用户提问 -> POST /api/chat/stream -> '
    'VectorStoreIndex 检索 (top_k=10) -> ModelScopeReranker 重排序 (top_n=5) -> '
    'DeepSeek LLM 生成 -> SSE 流式响应'
)
doc.add_paragraph(
    '推理模式 (USE_REASONING=true)：知识库检索 + 用户文件 FAISS 检索 -> '
    '合并上下文 -> httpx 直连 DeepSeek Reasoner (stream: true) -> '
    '真流式 SSE（reasoning_content + content 双通道逐 token 到达）'
)

doc.add_heading('3.2 用户交互路径', level=2)
doc.add_paragraph('系统提供三条独立的用户交互路径：')
table2 = doc.add_table(rows=4, cols=4, style='Light Grid Accent 1')
headers = ['路径', '端点', '模式', '说明']
data2 = [
    ['RAG 流式（推理）', 'POST /api/chat/stream', 'SSE 真流式',
     'USE_REASONING=true: DeepSeek Reasoner 原生思维链，reasoning+token 双通道流式'],
    ['RAG 非流式', 'POST /api/chat', 'JSON 响应',
     '非流式便捷接口，返回 answer + sources'],
    ['CrewAI Agent', 'POST /api/chat/agent', 'SSE 流式',
     'RAG 知识库 + 百度搜索 + Tavily 搜索三工具协作'],
]
for j, h in enumerate(headers):
    table2.cell(0, j).text = h
for i, row in enumerate(data2):
    for j, val in enumerate(row):
        table2.cell(i + 1, j).text = val
for row in table2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'
doc.add_paragraph()

doc.add_heading('3.3 索引单例架构', level=2)
doc.add_paragraph(
    '向量索引通过 index_singleton.py 管理，使用双重检查锁定（threading.Lock）保证线程安全。'
    '该单例被两条路径共享：1) RAG API 路由（routes.py），2) CrewAI RAG 工具（tools/rag_tool.py）。'
    '修改索引加载逻辑时必须保持双重检查锁定模式（if not -> with lock -> if not），'
    '否则高并发下可能多次重建或出现竞态条件。reset_index() 仅用于测试。'
)
doc.add_paragraph(
    '索引持久化时自动保存 embedding_config.json（含 dim/mode/model），'
    '下次加载时校验配置是否匹配，不匹配则拒绝加载并提示使用 --rebuild 重建。'
)

doc.add_heading('3.4 MCP 服务器与 Agent 架构', level=2)
doc.add_paragraph(
    'MCP 服务器（src/mcp_server/server.py）基于 FastMCP 框架，通过 @mcp.tool() 装饰器'
    '注册 baidu_search 和 tavily_search 两个工具，stdio 模式运行，供 MCP 客户端调用。'
)
doc.add_paragraph(
    'CrewAI Agent 管线：agent_routes.py 使用 asyncio.to_thread() 将同步 crew.kickoff() '
    '放入线程池执行。Agent 配备三个工具：RAGKnowledgeBaseTool（LlamaIndex 知识库检索）、'
    'MCPBaiduSearchTool（百度网页搜索）、MCPTavilySearchTool（Tavily 搜索引擎）。'
    '后两者在每次调用时启动 MCP stdio 子进程，内部使用 asyncio.run() 桥接同步/异步。'
)

# ====== 4. 项目结构 ======
doc.add_heading('4. 项目结构', level=1)

doc.add_heading('4.1 后端模块详表', level=2)
modules = [
    ('src/main.py', 'FastAPI 应用入口，CORS 配置，挂载 router/files_router/agent_router，启动时 validate_config() 校验 API Key；APP_ENV=development 时启用 /docs'),
    ('src/config.py', '所有环境变量和 RAG 参数，从项目根目录 .env 读取；含 validate_config()；导入即自动创建索引/上传目录'),
    ('src/api/routes.py', 'GET /api/health、POST /api/chat、POST /api/chat/stream；通过 index_singleton.get_index() 获取索引；USE_REASONING 分支选择推理模式或标准模式'),
    ('src/api/agent_routes.py', 'POST /api/chat/agent — CrewAI Agent 端点：组装 RAG 知识库 + 百度搜索 + Tavily 搜索三工具 + Engineer Agent + Task，SSE 流式输出'),
    ('src/api/files.py', 'POST /api/files/upload、GET /api/files、DELETE /api/files/{id}、POST /api/files/search'),
    ('src/api/schemas.py', 'Pydantic 请求/响应模型'),
    ('src/rag/index_singleton.py', '共享索引单例：双重检查锁定 + get_index() / reset_index()。被 RAG API 和 CrewAI RAG 工具共用'),
    ('src/rag/embedding.py', '双模式：LocalEmbedding (sentence-transformers 离线) 和 ModelScopeEmbedding (API，Qwen3-Embedding-8B)；异步方法使用 loop.run_in_executor 避免阻塞'),
    ('src/rag/reranker.py', 'ModelScopeReranker 封装 Rerank API；失败时打印 [WARN] 并降级为原始排序'),
    ('src/rag/indexing.py', '文档加载（统一解析器支持 6 种格式）、SentenceSplitter 中文切分、IngestionPipeline 处理、索引构建与持久化、embedding_config.json 元数据保存'),
    ('src/rag/query.py', 'get_query_engine() 非流式、get_streaming_query_engine() 流式、query_with_sources() 便捷封装、astream_query_structured() JSON 伪流式、astream_query_reasoning() Reasoner 真流式'),
    ('src/rag/prompts.py', '标准模式 (CONSTRUCTION_SYSTEM_PROMPT + CONSTRUCTION_QA_PROMPT) 和 JSON 结构化模式 (CONSTRUCTION_JSON_SYSTEM_PROMPT + CONSTRUCTION_JSON_QA_TMPL) 两套提示词'),
    ('src/rag/file_parser.py', '多格式解析：PDF (PyMuPDF)、Word (python-docx)、PPT (python-pptx)、Excel (openpyxl)、TXT/Markdown；统一入口 parse_file(file_path, file_type)；PDF 在 try/finally 中显式 close()'),
    ('src/rag/user_index.py', 'UserFAISSIndex 类：管理用户上传文件的 FAISS 索引（增删查 + 持久化），使用 IndexFlatIP + IndexIDMap 支持按 ID 删除'),
    ('src/rag/llm.py', 'get_llm() 工厂函数，通过 OpenAILike 接入 DeepSeek-Chat。注意：与 src/llm/model.py 同名不同实现，前者供 LlamaIndex RAG，后者供 CrewAI'),
    ('src/mcp_server/server.py', 'MCP 服务器，基于 FastMCP 框架，@mcp.tool() 注册 baidu_search 和 tavily_search 工具，stdio 模式运行'),
    ('src/mcp_server/baidu_tool.py', 'search_baidu() 包装函数，供 MCP 服务器调用'),
    ('src/mcp_server/tavily_tool.py', 'Tavily Search API 封装，调用 TavilyClient 返回搜索结果'),
    ('src/mcp_server/baidu_tools.py', '百度网页搜索爬虫：HTML 解析（BeautifulSoup + 自定义 HTMLParser 双解析器）、重定向跟踪、分页、反爬处理'),
    ('src/crew/agents.py', 'create_engineer_agent(llm, tools)：土木工程专家 agent 定义（role/goal/backstory）'),
    ('src/crew/tasks.py', 'create_engineering_task(agent, question)：结构化输出任务模板（原因 -> 方案 -> 专业语言）'),
    ('src/crew/crew.py', 'run_crew(agent, task)：Crew 编排器，组装 agent + task 并 kickoff()'),
    ('src/llm/model.py', 'DeepSeek LLM 工厂（供 CrewAI 使用），通过 crewai.LLM 包装，硬编码 deepseek-chat 模型'),
    ('src/tools/mcp_tools.py', 'MCPBaiduSearchTool + MCPTavilySearchTool — CrewAI BaseTool 子类，每次调用启动 MCP stdio 子进程，asyncio.run() 桥接同步/异步'),
    ('src/tools/rag_tool.py', 'RAGKnowledgeBaseTool — CrewAI BaseTool 子类，通过 index_singleton.get_index() 获取共享索引，query_with_sources() 检索知识库'),
    ('src/agent/agent.py', '早期 Agent 工厂原型 (create_agents())，当前未被使用，保留作参考'),
    ('src/agent/mcp_client.py', '异步 MCP 客户端演示脚本，通过 stdio 连接 MCP 服务器、列出工具、调用示例'),
    ('src/shared/', 'common.py (SearchResult/Evidence 数据类)、config.py (百度搜索参数)、crawler.py (Article 抓取器占位)'),
]
for name, desc in modules:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    p.add_run(desc).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

doc.add_heading('4.2 前端组件详表', level=2)
frontend_comps = [
    ('src/App.jsx', '根组件：Header（状态指示灯）+ FilePanel 侧边栏 + ChatWindow 主区域'),
    ('src/components/ChatWindow.jsx', '对话主体：消息列表、SSE 流式渲染（reasoning/analysis/token/source/tool_call/done/error 事件）、建议问题、输入框、localStorage 缓存持久化'),
    ('src/components/MessageBubble.jsx', '消息气泡：marked 库渲染 GFM -> DOMPurify.sanitize() 消毒、流式思考面板、分析摘要折叠面板、来源引用卡片'),
    ('src/components/SourceCard.jsx', '引用卡片：支持标准规范引用和用户文件引用两种样式'),
    ('src/components/FilePanel.jsx', '文件管理侧边栏：可折叠、文件列表（含图标/大小/时间）、删除确认'),
    ('src/components/FileUploadZone.jsx', '文件上传区：拖拽/点击上传、多状态（default/dragging/uploading/success/error）、客户端校验'),
    ('src/api/chat.js', 'sendMessage() 非流式、sendMessageStream() SSE 流式（AsyncGenerator + AbortController 取消）、sendAgentMessageStream() Agent 流式、healthCheck()'),
    ('src/api/files.js', 'uploadFile()、listFiles()、deleteFile()'),
    ('src/App.css', '全局样式 (~1078 行)：Design tokens (CSS 变量 Slate 色系)、布局、所有组件样式、响应式断点、动画'),
]
for name, desc in frontend_comps:
    p = doc.add_paragraph()
    run = p.add_run(f'{name}：')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    p.add_run(desc).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(2)

# ====== 5. 环境变量 ======
doc.add_heading('5. 环境变量与配置', level=1)
doc.add_paragraph('所有配置在项目根目录 .env 文件中（模板见 .env.example）。以下为主要环境变量：')

env_table = doc.add_table(rows=16, cols=3, style='Light Grid Accent 1')
env_data = [
    ('变量', '默认值', '说明'),
    ('DEEPSEEK_API_KEY', '—', '【必填】DeepSeek API 密钥'),
    ('DEEPSEEK_BASE_URL', 'https://api.deepseek.com', 'DeepSeek API 地址'),
    ('DEEPSEEK_MODEL', 'deepseek-chat', '标准对话模型'),
    ('DEEPSEEK_REASONING_MODEL', 'deepseek-reasoner', '推理模型，USE_REASONING=true 时使用'),
    ('USE_REASONING', 'false', 'true 时启用 DeepSeek Reasoner 真流式推理模式'),
    ('MODELSCOPE_API_KEY', '—', 'ModelScope API 密钥（API Embedding/Rerank 需要）'),
    ('TAVILY_API_KEY', '—', 'Tavily Search API 密钥（Agent 搜索需要）'),
    ('MODELSCOPE_BASE_URL', 'https://api.modelscope.cn', 'ModelScope API 地址'),
    ('EMBEDDING_MODE', 'local', 'local（离线，BAAI/bge-small-zh-v1.5）或 api（云端，Qwen3-Embedding-8B）'),
    ('LOCAL_EMBEDDING_MODEL', 'BAAI/bge-small-zh-v1.5', '本地 Embedding 模型名（~100MB，中文优化）'),
    ('EMBEDDING_MODEL', 'Qwen/Qwen3-Embedding-8B', 'API Embedding 模型（仅 EMBEDDING_MODE=api 时使用）'),
    ('RERANK_MODEL', 'Qwen/Qwen3-Reranker-8B', 'Rerank 模型（ModelScope API）'),
    ('BACKEND_PORT', '8000', '后端服务端口'),
    ('FRONTEND_PORT', '3000', '前端服务端口（修改后需同步 CORS 和 vite.config.js proxy）'),
    ('APP_ENV', 'development', '运行环境，影响 Swagger 可见性 (/docs) 和热重载'),
]
for i, (col1, col2, col3) in enumerate(env_data):
    env_table.cell(i, 0).text = col1
    env_table.cell(i, 1).text = col2
    env_table.cell(i, 2).text = col3
    for cell in env_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'

doc.add_paragraph()
doc.add_paragraph('关键 RAG 参数（在 config.py 中固定）：')
rag_params = [
    'CHUNK_SIZE = 512（文本切分大小）',
    'CHUNK_OVERLAP = 64（切分重叠量）',
    'TOP_K_RETRIEVE = 10（初次检索数量）',
    'TOP_K_RERANK = 5（重排序后保留数量）',
    'USER_TOP_K = 5（用户文件检索数量）',
    'EMBEDDING_DIM = 512（向量维度，必须与 Embedding 模型输出一致）',
    'MAX_FILE_SIZE = 50MB',
    'ALLOWED_EXTENSIONS = .pdf, .docx, .pptx, .xlsx, .txt, .md',
]
for param in rag_params:
    doc.add_paragraph(f'  - {param}')

doc.add_paragraph()
doc.add_paragraph('存储路径说明：')
doc.add_paragraph('  - 标准知识库索引：backend/storage/')
doc.add_paragraph('  - 用户 FAISS 索引：backend/storage/user_faiss/')
doc.add_paragraph('  - 用户上传文件：backend/uploads/')
doc.add_paragraph('  - 知识库文档：backend/src/data/documents/（支持递归子目录，按 国家标准GB/行业标准/地方标准/技术规程/法律法规 分类）')

# ====== 6. API 参考 ======
doc.add_heading('6. API 参考', level=1)

doc.add_heading('6.1 健康检查', level=2)
doc.add_paragraph('端点：GET /api/health')
doc.add_paragraph('返回：')
doc.add_paragraph('  {"llm_model": "deepseek-chat", "embedding_model": "Qwen/Qwen3-Embedding-8B", '
                  '"rerank_model": "Qwen/Qwen3-Reranker-8B", "index_ready": true, "user_files_count": 0}')

doc.add_heading('6.2 RAG 问答', level=2)
doc.add_paragraph('非流式 — POST /api/chat')
doc.add_paragraph('  请求：{"question": "梁钢筋搭接长度是多少？"}')
doc.add_paragraph('  返回：{"answer": "...", "sources": [...], "question": "..."}')
doc.add_paragraph()
doc.add_paragraph('流式（推理）— POST /api/chat/stream')
doc.add_paragraph('  USE_REASONING=true 时使用 DeepSeek Reasoner 真流式推理')
doc.add_paragraph('  SSE 事件类型：reasoning（思考过程流式）| token（回答流式）| source（来源）| done | error')
doc.add_paragraph('  USE_REASONING=false 时使用标准 LlamaIndex 流式查询引擎')

doc.add_heading('6.3 Agent 问答', level=2)
doc.add_paragraph('端点：POST /api/chat/agent')
doc.add_paragraph('说明：CrewAI Agent 管线，RAG 知识库 + 百度搜索 + Tavily 搜索三工具协作')
doc.add_paragraph('SSE 事件类型：token（回答文本）| tool_call（工具调用记录）| done | error')

doc.add_heading('6.4 文件管理', level=2)
file_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
file_data = [
    ('方法', '端点', '说明'),
    ('POST', '/api/files/upload', '上传文档（支持 .pdf/.docx/.pptx/.xlsx/.txt/.md，最大 50MB）'),
    ('GET', '/api/files', '获取已上传文件列表'),
    ('GET', '/api/files/{id}', '获取单个文件元信息'),
    ('DELETE', '/api/files/{id}', '删除指定文件及其 FAISS 向量索引'),
    ('POST', '/api/files/search', '在用户文件中语义搜索'),
]
for i, (method, endpoint, desc) in enumerate(file_data):
    file_table.cell(i, 0).text = method
    file_table.cell(i, 1).text = endpoint
    file_table.cell(i, 2).text = desc
    for cell in file_table.rows[i].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = '微软雅黑'

# ====== 7. 安装与运行 ======
doc.add_heading('7. 安装与运行', level=1)

doc.add_heading('7.1 环境要求', level=2)
for req in ['Python 3.10+', 'Node.js 18+',
            '（可选）ModelScope API Key（使用云端 Embedding/Rerank 时需要）',
            '（可选）Tavily API Key（使用 Agent 搜索功能时需要）',
            '本地 Embedding 模式首次启动自动下载约 100MB 模型文件（BAAI/bge-small-zh-v1.5）']:
    doc.add_paragraph(f'  - {req}')

doc.add_heading('7.2 快速启动', level=2)
steps = [
    ('步骤 1：配置环境变量',
     'cp .env.example .env\n编辑 .env，填入 DEEPSEEK_API_KEY=sk-xxxxxxxx'),
    ('步骤 2：安装后端依赖',
     'cd backend\npip install -r requirements.txt\n'
     '# 可选的 Agent/MCP 依赖：\n'
     'pip install mcp fastmcp tavily-python beautifulsoup4 crewai'),
    ('步骤 3：安装前端依赖',
     'cd frontend\nnpm install'),
    ('步骤 4：构建知识库索引',
     'cd backend\npython scripts/build_index.py              # 构建\n'
     'python scripts/build_index.py --rebuild    # 强制重建\n'
     'python scripts/build_index.py --status     # 查看状态'),
    ('步骤 5：启动服务',
     '终端1（后端）：cd backend && python -m src.main  -> http://localhost:8000\n'
     '终端2（前端）：cd frontend && npm run dev  -> http://localhost:3000'),
    ('步骤 6：使用',
     '访问 http://localhost:3000 开始提问'),
]
for title_text, cmd in steps:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.font.bold = True
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    doc.add_paragraph(f'  {cmd}')

doc.add_heading('7.3 数据管道', level=2)
doc.add_paragraph('以下脚本用于扩充知识库文档：')
pipelines = [
    ('python scripts/02_setup_knowledge_base.py',
     '初始化知识库目录结构 + PDF 文本提取（PyMuPDF） + OCR 支持（--convert-pdf --ocr） + 知识库统计（--stats）'),
    ('python scripts/03_auto_crawler.py',
     '从 gf.cabr-fire.com 自动爬取标准全文：搜索 -> 匹配 -> 提取章节 -> 合并保存。'
     '  --batch --priority S,A --missing-only 批量爬取高优先级标准'),
    ('python scripts/01_download.py',
     '从官方来源下载规范（flk.npc.gov.cn / openstd.samr.gov.cn），--type law 下载法律法规'),
]
for cmd, desc in pipelines:
    p = doc.add_paragraph()
    run = p.add_run(cmd)
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    doc.add_paragraph(f'  {desc}')

# ====== 8. 关键实现细节 ======
doc.add_heading('8. 关键实现细节', level=1)
doc.add_paragraph('以下为实现中非显而易见的约束，违反会导致隐蔽 bug：')

details = [
    ('FAISS 向量必须 L2 归一化',
     'UserFAISSIndex 使用 IndexFlatIP（内积），只有向量经过 L2 归一化后，内积才等价于余弦相似度。'
     '构建索引时存储前和查询时 encode 都必须传入 normalize_embeddings=True。'
     '如果忘记归一化，检索结果会严重偏离——内积同时受方向和模长影响，长文本天然得分更高。'),
    ('LocalEmbedding 异步方法不能阻塞事件循环',
     'sentence-transformers 的 encode() 是同步 CPU 密集型调用，在 async def 中直接调用会阻塞整个事件循环。'
     'embedding.py 中的 _aget_* 方法必须使用 loop.run_in_executor(None, lambda: self._model.encode(...))。'),
    ('索引单例的双重检查锁定',
     'index_singleton.py 的 get_index() 使用 if->with lock->if 双重检查模式，'
     '防止多个并发请求同时触发索引重建。该单例被 RAG API 和 CrewAI RAG 工具共享。'),
    ('CrewAI Agent 同步/异步桥接',
     'agent_routes.py 使用 asyncio.to_thread(crew.kickoff) 放入线程池执行，避免阻塞 FastAPI 事件循环。'
     'MCPBaiduSearchTool/MCPTavilySearchTool 在同步 _run() 中调用 asyncio.run() 桥接 MCP 异步协议，'
     '每次工具调用都启动新的 MCP stdio 子进程（短生命周期），并设置 PYTHONPATH 和 PYTHONIOENCODING=utf-8。'),
    ('embedding_config.json 校验',
     '构建索引时保存当前 embedding 配置（dim/mode/model）；加载索引时自动对比，'
     '维度或模型不匹配则抛出 ValueError 提示使用 --rebuild 重建，防止 512 维 / 1024 维向量混用。'),
    ('XSS 防护：Markdown 输出必须消毒',
     'MessageBubble.jsx 使用 DOMPurify.sanitize() 消毒 marked 库渲染的 HTML，'
     '防止 LLM 输出中的恶意脚本。修改前端渲染逻辑时不能跳过此步骤。'),
    ('PDF 文件句柄必须释放',
     'file_parser.py 中 PyMuPDF (fitz.open()) 返回的 Document 对象必须在 try/finally 中显式调用 .close()，'
     '否则会导致文件句柄泄漏，长时间运行后可能耗尽文件描述符。'),
    ('SSE 流支持取消',
     '前端 chat.js 的 sendMessageStream() 接受 options.signal（AbortController），'
     '用于取消正在进行的 SSE 流。修改流式请求逻辑时应保持此支持，否则用户切换页面后流仍占用后端连接。'),
    ('Reranker 降级机制',
     'ModelScope Rerank API 失败时自动回退到原始检索排序，打印 [WARN] 日志但不阻塞请求。'
     '该机制确保在网络不稳定或 API Key 未配置时系统仍然可用。'),
    ('文件解析惰性导入',
     'file_parser.py 中各格式解析库在函数内部惰性导入（import fitz / from docx import Document 等），'
     '按需加载，减少启动开销。'),
    ('错误消息脱敏',
     'API 返回给客户端的错误消息不应包含内部路径、堆栈跟踪或敏感配置信息。'),
]
for title_text, detail_text in details:
    p = doc.add_paragraph()
    run = p.add_run(f'{title_text}：')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '微软雅黑'
    p2 = doc.add_paragraph(detail_text)
    p2.paragraph_format.space_after = Pt(6)
    p2.style.font.size = Pt(10)

# ====== 9. 架构约束 ======
doc.add_heading('9. 架构约束', level=1)
constraints = [
    '无数据库 — 知识库全部为 .txt/.pdf/.docx 等文件格式，标准向量索引 + 用户 FAISS 索引均持久化为文件（backend/storage/），无 PostgreSQL/Redis。',
    '无用户认证 — 无登录、无用户管理。API Key（DeepSeek/ModelScope/Tavily）在服务启动时由 validate_config() 校验。',
    '单用户设计 — 索引通过 index_singleton.py 全局单例管理，files.py 中全局 _user_index 单例，无会话管理，对话历史仅在浏览器端 localStorage 持久化。',
    '无自动化测试 — 项目当前无 pytest/Vitest 测试文件。test_tools.py 是 MCP 工具列表验证脚本，test_crew.py 是 Agent 手动测试脚本。修改代码后需手动启动服务验证。',
    '索引懒加载 — 首次 API 调用或 Agent 工具调用时才构建/加载索引，非启动时加载。',
    'Reranker 降级 + 日志 — ModelScope Rerank API 失败时自动回退到原始检索排序，同时打印 [WARN] 日志提示检查 API Key 和网络。',
    '文件解析惰性导入 — file_parser.py 中各格式解析库在函数内部惰性导入，按需加载。',
    'Swagger 条件启用 — APP_ENV=development 时启用 /docs，非开发环境隐藏。',
    'CORS 硬编码 — main.py 仅允许 localhost:3000 和 127.0.0.1:3000。修改 FRONTEND_PORT 时需同步更新 CORS origins 和 vite.config.js 的 proxy target。',
    'EMBEDDING_DIM 固定 — FAISS 向量维度必须与 Embedding 模型输出维度一致。更换 Embedding 模型时需同步修改 config.py 中的 EMBEDDING_DIM 并重建索引（--rebuild）。',
    '用户文件仅在 USE_REASONING 模式生效 — 标准流式模式（USE_REASONING=false）不检索用户上传文件，仅检索标准知识库。',
    'MCP/CrewAI 依赖未纳入 requirements.txt — 以下依赖需手动安装：mcp、fastmcp、tavily-python、beautifulsoup4、crewai。',
    'config.py 导入即创建目录 — INDEX_STORAGE_DIR、DOCUMENTS_DIR、USER_UPLOADS_DIR、USER_FAISS_DIR 在模块导入时自动 mkdir()。',
    'src/rag/llm.py 与 src/llm/model.py 同名不同实现 — 前者供 LlamaIndex RAG 使用（OpenAILike），后者供 CrewAI 使用（crewai.LLM），修改时注意区分。',
]
for c in constraints:
    doc.add_paragraph(f'  - {c}')

# ====== 10. 部署说明 ======
doc.add_heading('10. 部署说明', level=1)

doc.add_heading('开发环境', level=2)
doc.add_paragraph('  后端：python -m src.main（uvicorn + 热重载，端口 8000）')
doc.add_paragraph('  前端：npm run dev（Vite 开发服务器，端口 3000，自动代理 /api -> 后端）')

doc.add_heading('生产环境', level=2)
doc.add_paragraph('  前端构建：cd frontend && npm run build -> 输出 dist/ 静态文件')
doc.add_paragraph('  部署 dist/ 到 Nginx / CDN，后端 API 通过反向代理或直接暴露')
doc.add_paragraph('  后端生产启动：uvicorn src.main:app --host 0.0.0.0 --port 8000')
doc.add_paragraph('  APP_ENV=production（关闭热重载，隐藏 /docs）')

doc.add_heading('部署前 Checklist', level=2)
checklist = [
    '确认 .env 中 DEEPSEEK_API_KEY 已配置',
    '确认 APP_ENV=production（隐藏 /docs Swagger）',
    '更新 main.py 中 CORS allow_origins 以匹配前端实际域名',
    '确认知识库索引已构建完毕（python scripts/build_index.py）',
    '确认 EMBEDDING_MODE=local 或已配置 MODELSCOPE_API_KEY',
    '如使用 Agent 功能，确认 TAVILY_API_KEY 已配置',
    '确认 FRONTEND_PORT 配置与前端实际访问端口一致',
]
for item in checklist:
    doc.add_paragraph(f'  - {item}')

# ====== 保存 ======
output_path = r'D:\Desktop\my project mcp\项目1技术文档.docx'
doc.save(output_path)
print(f'文档已生成：{output_path}')
